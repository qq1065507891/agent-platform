from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Iterable, Literal
import hashlib
import os
from uuid import uuid4
from urllib.parse import urlparse
import logging

from fastapi import UploadFile
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument


def _disable_chroma_telemetry() -> None:
    os.environ["ANONYMIZED_TELEMETRY"] = "False"
    os.environ["CHROMA_TELEMETRY"] = "False"
    os.environ["CHROMADB_ANONYMIZED_TELEMETRY"] = "False"
    os.environ["POSTHOG_DISABLED"] = "true"

    if os.environ.get("CHROMA_PRODUCT_TELEMETRY_IMPL", "").strip().lower() == "none":
        os.environ.pop("CHROMA_PRODUCT_TELEMETRY_IMPL", None)

    try:
        import posthog  # type: ignore

        posthog.disabled = True

        def _noop_capture(*_args, **_kwargs):
            return None

        posthog.capture = _noop_capture  # type: ignore[attr-defined]
    except Exception:
        pass


_disable_chroma_telemetry()

import chromadb
from chromadb.config import Settings as ChromaSettings

logger = logging.getLogger(__name__)
logger.info(
    "[chroma-telemetry] anonymized=%s chroma=%s chromadb=%s posthog_disabled=%s telemetry_impl=%s",
    os.environ.get("ANONYMIZED_TELEMETRY"),
    os.environ.get("CHROMA_TELEMETRY"),
    os.environ.get("CHROMADB_ANONYMIZED_TELEMETRY"),
    os.environ.get("POSTHOG_DISABLED"),
    os.environ.get("CHROMA_PRODUCT_TELEMETRY_IMPL", "<unset>"),
)

from app.core.config import settings
from app.services.embeddings import get_embeddings


@dataclass
class IngestResult:
    doc_id: str
    version: int
    chunk_count: int
    ingest_status: str = "indexed"


@dataclass
class KnowledgeDocumentItem:
    doc_id: str
    source: str
    version: int
    chunk_count: int
    status: str
    created_at: str


class RAGIngestError(Exception):
    pass


class RAGUnsupportedFileTypeError(RAGIngestError):
    pass


class RAGTextExtractionError(RAGIngestError):
    pass


class RAGService:
    _client_cache: chromadb.ClientAPI | None = None
    _embeddings_cache: Embeddings | None = None
    _store_cache: Chroma | None = None

    def __init__(self) -> None:
        self._persist_path = Path(settings.chroma_persist_path).resolve()
        self._collection_name = "agent_knowledge"

    def _client(self) -> chromadb.ClientAPI:
        if RAGService._client_cache is not None:
            return RAGService._client_cache

        chroma_settings = ChromaSettings(anonymized_telemetry=False)
        chroma_url_raw = str(getattr(settings, "chroma_url", "") or "")
        chroma_url = chroma_url_raw.strip()
        if " #" in chroma_url:
            chroma_url = chroma_url.split(" #", 1)[0].strip()
        if chroma_url.startswith("#"):
            chroma_url = ""

        if chroma_url:
            parsed = urlparse(chroma_url if "://" in chroma_url else f"http://{chroma_url}")
            host = parsed.hostname or chroma_url
            port = parsed.port
            ssl = parsed.scheme == "https"

            RAGService._client_cache = chromadb.HttpClient(
                host=host,
                port=port,
                ssl=ssl,
                settings=chroma_settings,
            )
            logger.info("[rag-chroma] using remote chroma host=%s port=%s ssl=%s", host, port, ssl)
        else:
            self._persist_path.mkdir(parents=True, exist_ok=True)
            RAGService._client_cache = chromadb.PersistentClient(path=str(self._persist_path), settings=chroma_settings)
            logger.info("[rag-chroma] using local chroma persist_path=%s", self._persist_path)

        return RAGService._client_cache

    def _embedding_function(self) -> Embeddings:
        if RAGService._embeddings_cache is None:
            RAGService._embeddings_cache = get_embeddings()
        return RAGService._embeddings_cache

    def _vector_store(self) -> Chroma:
        if RAGService._store_cache is None:
            client = self._client()
            RAGService._store_cache = Chroma(
                client=client,
                collection_name=self._collection_name,
                embedding_function=self._embedding_function(),
            )
        return RAGService._store_cache

    def _split_text(self, text: str, *, doc_type: str) -> list[str]:
        chunk_size = int(getattr(settings, "rag_chunk_size", 800))
        chunk_overlap = int(getattr(settings, "rag_chunk_overlap", 120))

        if doc_type == "pdf":
            chunk_size = int(getattr(settings, "rag_pdf_chunk_size", chunk_size))
            chunk_overlap = int(getattr(settings, "rag_pdf_chunk_overlap", chunk_overlap))
        elif doc_type in {"doc", "docx"}:
            chunk_size = int(getattr(settings, "rag_docx_chunk_size", chunk_size))
            chunk_overlap = int(getattr(settings, "rag_docx_chunk_overlap", chunk_overlap))

        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        return splitter.split_text(text)

    def _extract_text_from_pdf(self, data: bytes) -> str:
        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _extract_text_from_docx(self, data: bytes) -> str:
        doc = DocxDocument(BytesIO(data))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    def _extract_text(self, filename: str, data: bytes) -> str:
        suffix = Path(filename).suffix.lower()

        try:
            if suffix == ".pdf":
                return self._extract_text_from_pdf(data)
            if suffix in {".doc", ".docx"}:
                return self._extract_text_from_docx(data)
            if suffix in {"", ".txt", ".md", ".csv", ".json", ".log"}:
                return data.decode("utf-8", errors="ignore")
        except Exception as exc:
            raise RAGTextExtractionError(f"failed to extract text from {suffix or 'unknown'}: {exc}") from exc

        raise RAGUnsupportedFileTypeError(f"unsupported file type: {suffix or '<none>'}")

    def _infer_doc_type(self, filename: str) -> str:
        suffix = Path(filename or "uploaded").suffix.lower()
        return suffix.lstrip(".") or "txt"

    def _compute_content_hash(self, data: bytes) -> str:
        return hashlib.sha256(data).hexdigest()

    def _estimate_tokens(self, text: str) -> int:
        return max(1, len(text) // 4)

    def _exists_by_hash(self, *, content_hash: str, agent_id: str | None = None) -> bool:
        store = self._vector_store()
        try:
            filter_payload: dict[str, str] = {"content_hash": content_hash}
            if agent_id:
                filter_payload["agent_id"] = agent_id

            result = store.get(where=filter_payload)
            ids = result.get("ids", []) if isinstance(result, dict) else []
            return bool(ids)
        except Exception:
            return False

    def _build_documents(
        self,
        chunks: Iterable[str],
        doc_id: str,
        source_name: str,
        *,
        doc_type: str,
        content_hash: str,
    ) -> list[Document]:
        timestamp = datetime.now(timezone.utc).isoformat()
        built_documents: list[Document] = []
        for index, chunk in enumerate(chunks):
            built_documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "doc_id": doc_id,
                        "source": source_name,
                        "created_at": timestamp,
                        "chunk_index": index,
                        "doc_type": doc_type,
                        "token_estimate": self._estimate_tokens(chunk),
                        "content_hash": content_hash,
                    },
                )
            )
        return built_documents

    async def ingest_upload(self, file: UploadFile, agent_id: str | None = None) -> IngestResult:
        data = await file.read()
        doc_id = f"doc-{uuid4().hex}"
        version = 1
        filename = file.filename or "uploaded"
        doc_type = self._infer_doc_type(filename)
        content_hash = self._compute_content_hash(data)

        dedup_enabled = bool(getattr(settings, "rag_ingest_dedup_enabled", True))
        if dedup_enabled and self._exists_by_hash(content_hash=content_hash, agent_id=agent_id):
            return IngestResult(doc_id=doc_id, version=version, chunk_count=0, ingest_status="deduplicated")

        content = self._extract_text(filename, data)
        chunks = [chunk.strip() for chunk in self._split_text(content, doc_type=doc_type) if chunk.strip()]
        documents = self._build_documents(
            chunks,
            doc_id,
            filename,
            doc_type=doc_type,
            content_hash=content_hash,
        )
        if agent_id:
            for doc in documents:
                doc.metadata["agent_id"] = agent_id
        for doc in documents:
            doc.metadata["version"] = version
            doc.metadata["status"] = "indexed"

        if not documents:
            return IngestResult(doc_id=doc_id, version=version, chunk_count=0, ingest_status="empty")

        store = self._vector_store()
        store.add_documents(documents)
        logger.info(
            "[rag-ingest] doc_id=%s source=%s doc_type=%s chunks=%s agent_id=%s",
            doc_id,
            filename,
            doc_type,
            len(documents),
            agent_id,
        )
        return IngestResult(doc_id=doc_id, version=version, chunk_count=len(documents), ingest_status="indexed")

    def list_agent_documents(self, *, agent_id: str, keyword: str | None = None) -> list[KnowledgeDocumentItem]:
        store = self._vector_store()
        result = store.get(where={"agent_id": agent_id}, include=["metadatas"])
        metadatas = result.get("metadatas", []) if isinstance(result, dict) else []

        docs_map: dict[str, KnowledgeDocumentItem] = {}
        for metadata in metadatas:
            if not isinstance(metadata, dict):
                continue
            doc_id = str(metadata.get("doc_id") or "").strip()
            if not doc_id:
                continue

            source = str(metadata.get("source") or "unknown")
            status = str(metadata.get("status") or "indexed")
            version = int(metadata.get("version") or 1)
            created_at = str(metadata.get("created_at") or "")

            current = docs_map.get(doc_id)
            if current is None:
                docs_map[doc_id] = KnowledgeDocumentItem(
                    doc_id=doc_id,
                    source=source,
                    version=version,
                    chunk_count=1,
                    status=status,
                    created_at=created_at,
                )
            else:
                current.chunk_count += 1
                if created_at and (not current.created_at or created_at > current.created_at):
                    current.created_at = created_at

        items = list(docs_map.values())
        keyword_value = (keyword or "").strip().lower()
        if keyword_value:
            items = [item for item in items if keyword_value in item.doc_id.lower() or keyword_value in item.source.lower()]

        items.sort(key=lambda item: item.created_at, reverse=True)
        return items

    def _build_delete_where_payload(self, *, agent_id: str, doc_id: str, delete_mode: Literal["soft", "hard"]) -> dict:
        if delete_mode == "soft":
            return {"$and": [{"agent_id": agent_id}, {"doc_id": doc_id}, {"status": "indexed"}]}
        return {"agent_id": agent_id, "doc_id": doc_id}

    def delete_agent_document(self, *, agent_id: str, doc_id: str, delete_mode: Literal["soft", "hard"] = "soft") -> int:
        store = self._vector_store()
        where_payload = self._build_delete_where_payload(agent_id=agent_id, doc_id=doc_id, delete_mode=delete_mode)
        existing = store.get(where=where_payload)
        ids = existing.get("ids", []) if isinstance(existing, dict) else []
        if not ids:
            return 0

        if delete_mode == "hard":
            store.delete(where=where_payload)
            return len(ids)

        metadatas = existing.get("metadatas", []) if isinstance(existing, dict) else []
        documents = existing.get("documents", []) if isinstance(existing, dict) else []
        if not isinstance(metadatas, list):
            metadatas = []
        if not isinstance(documents, list):
            documents = []

        rewritten_documents: list[Document] = []
        for idx, _ in enumerate(ids):
            metadata = metadatas[idx] if idx < len(metadatas) and isinstance(metadatas[idx], dict) else {}
            content = str(documents[idx] if idx < len(documents) else "")
            rewritten_meta = {**metadata, "status": "deleted"}
            rewritten_documents.append(Document(page_content=content, metadata=rewritten_meta))

        store.delete(ids=ids)
        if rewritten_documents:
            store.add_documents(rewritten_documents, ids=ids)
        return len(ids)

    def delete_agent_documents(
        self,
        *,
        agent_id: str,
        doc_ids: list[str],
        delete_mode: Literal["soft", "hard"] = "soft",
    ) -> dict[str, int]:
        results: dict[str, int] = {}
        for doc_id in doc_ids:
            normalized = str(doc_id or "").strip()
            if not normalized:
                continue
            results[normalized] = self.delete_agent_document(
                agent_id=agent_id,
                doc_id=normalized,
                delete_mode=delete_mode,
            )
        return results

    def purge_deleted_documents(self, *, agent_id: str | None = None, doc_ids: list[str] | None = None) -> dict[str, int]:
        store = self._vector_store()

        if doc_ids:
            results: dict[str, int] = {}
            for doc_id in doc_ids:
                normalized = str(doc_id or "").strip()
                if not normalized:
                    continue
                if agent_id:
                    where_payload = {"$and": [{"agent_id": agent_id}, {"doc_id": normalized}, {"status": "deleted"}]}
                else:
                    where_payload = {"$and": [{"doc_id": normalized}, {"status": "deleted"}]}
                existing = store.get(where=where_payload)
                ids = existing.get("ids", []) if isinstance(existing, dict) else []
                if ids:
                    store.delete(where=where_payload)
                results[normalized] = len(ids)
            return results

        where_payload = {"status": "deleted"}
        if agent_id:
            where_payload = {"$and": [{"agent_id": agent_id}, {"status": "deleted"}]}
        existing = store.get(where=where_payload)
        ids = existing.get("ids", []) if isinstance(existing, dict) else []
        if ids:
            store.delete(where=where_payload)

        return {"__all__": len(ids)}

    def as_retriever(self, agent_id: str | None = None):
        store = self._vector_store()
        recall_k = int(getattr(settings, "rag_recall_k", 24))
        fetch_k = int(getattr(settings, "rag_mmr_fetch_k", max(recall_k * 2, 32)))
        lambda_mult = float(getattr(settings, "rag_mmr_lambda_mult", 0.5))
        search_type = str(getattr(settings, "rag_search_type", "mmr") or "mmr").lower()

        search_kwargs: dict[str, object] = {"k": recall_k}
        if agent_id:
            search_kwargs["filter"] = {"$and": [{"agent_id": agent_id}, {"status": "indexed"}]}

        if search_type == "mmr":
            search_kwargs.update({"fetch_k": fetch_k, "lambda_mult": lambda_mult})
            return store.as_retriever(search_type="mmr", search_kwargs=search_kwargs)

        return store.as_retriever(search_kwargs=search_kwargs)

    def has_agent_knowledge(self, agent_id: str | None) -> bool:
        if not agent_id:
            return False

        try:
            store = self._vector_store()
            result = store.get(where={"$and": [{"agent_id": agent_id}, {"status": "indexed"}]}, limit=1)
            ids = result.get("ids", []) if isinstance(result, dict) else []
            return bool(ids)
        except Exception:
            return False

    def retrieve(self, query: str, agent_id: str | None = None) -> list[Document]:
        min_score = float(getattr(settings, "rag_similarity_min_score", 0.0))
        search_type = str(getattr(settings, "rag_search_type", "mmr") or "mmr").lower()
        diversify_enabled = bool(getattr(settings, "rag_threshold_diversify_enabled", True))

        if min_score <= 0:
            retriever = self.as_retriever(agent_id=agent_id)
            docs = retriever.invoke(query)
            logger.debug("[rag-retrieve] query_len=%s agent_id=%s hits=%s", len(query), agent_id, len(docs))
            return docs

        store = self._vector_store()
        filter_payload = {"$and": [{"agent_id": agent_id}, {"status": "indexed"}]} if agent_id else {"status": "indexed"}
        recall_k = max(int(getattr(settings, "rag_recall_k", 24)), 1)
        fetch_k = max(int(getattr(settings, "rag_mmr_fetch_k", max(recall_k * 2, 32))), recall_k)

        scored_docs = store.similarity_search_with_relevance_scores(
            query,
            k=fetch_k,
            filter=filter_payload,
        )

        thresholded = [(doc, score) for doc, score in scored_docs if score >= min_score]
        if not thresholded:
            logger.debug(
                "[rag-retrieve] query_len=%s agent_id=%s scored_hits=%s threshold_hits=0 min_score=%.3f => return empty",
                len(query),
                agent_id,
                len(scored_docs),
                min_score,
            )
            return []

        effective = thresholded

        if search_type == "mmr" and diversify_enabled:
            grouped: dict[str, list[tuple[Document, float]]] = defaultdict(list)
            for doc, score in effective:
                source = str(doc.metadata.get("source") or "")
                grouped[source].append((doc, score))

            for source_docs in grouped.values():
                source_docs.sort(key=lambda item: item[1], reverse=True)

            diversified: list[Document] = []
            keys = list(grouped.keys())
            cursor = 0
            while keys and len(diversified) < recall_k:
                key = keys[cursor]
                if grouped[key]:
                    doc, _score = grouped[key].pop(0)
                    diversified.append(doc)
                if not grouped[key]:
                    keys.pop(cursor)
                    if not keys:
                        break
                    cursor %= len(keys)
                else:
                    cursor = (cursor + 1) % len(keys)

            result_docs = diversified
        else:
            result_docs = [doc for doc, _score in effective[:recall_k]]

        threshold_hit_rate = (len(thresholded) / len(scored_docs)) if scored_docs else 0.0
        source_count = len({str(doc.metadata.get("source") or "") for doc in result_docs})
        source_diversity_ratio = (source_count / len(result_docs)) if result_docs else 0.0

        logger.debug(
            "[rag-retrieve] query_len=%s agent_id=%s scored_hits=%s threshold_hits=%s threshold_hit_rate=%.3f result_hits=%s source_diversity_ratio=%.3f min_score=%.3f search_type=%s diversify_enabled=%s",
            len(query),
            agent_id,
            len(scored_docs),
            len(thresholded),
            threshold_hit_rate,
            len(result_docs),
            source_diversity_ratio,
            min_score,
            search_type,
            diversify_enabled,
        )
        return result_docs

    def format_docs(self, docs: list[Document]) -> str:
        snippets = []
        for doc in docs:
            source = doc.metadata.get("source", "")
            doc_id = doc.metadata.get("doc_id", "")
            version = doc.metadata.get("version", "")
            snippets.append(f"[{doc_id}][v{version}][{source}] {doc.page_content}")
        return "\n\n".join(snippets)


def get_rag_service() -> RAGService:
    return RAGService()
